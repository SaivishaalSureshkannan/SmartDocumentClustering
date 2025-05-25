from typing import Union, List
import os
from fastapi import FastAPI, UploadFile, File, Body
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import aiofiles
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
import shutil
from data_store import document_store
from preprocessing import preprocess_text, tokens_to_string
from vectorization import document_vectorizer
from clustering import document_clusterer
from semantic_search import SemanticSearch
import numpy as np

app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Vite's default port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create uploads directory if it doesn't exist
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

async def save_uploaded_file(upload_file: UploadFile) -> str:
    """Save an uploaded file to the uploads directory"""
    file_path = os.path.join(UPLOAD_DIR, upload_file.filename)
    async with aiofiles.open(file_path, 'wb') as out_file:
        content = await upload_file.read()
        await out_file.write(content)
    return file_path

def extract_text(file_path: str) -> str:
    """Extract text from different file types"""
    _, ext = os.path.splitext(file_path.lower())
    
    try:
        if ext == '.pdf':
            # Extract text with better formatting
            text = extract_pdf_text(file_path)
            # Clean up the text
            lines = text.split('\n')
            # Remove empty lines and excessive whitespace
            lines = [line.strip() for line in lines if line.strip()]
            # Join with single newlines
            cleaned_text = '\n'.join(lines)
            # Replace form feed with page break marker
            cleaned_text = cleaned_text.replace('\f', '\n\n[PAGE BREAK]\n\n')
            return cleaned_text
            
        elif ext == '.docx':
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                # Handle different heading levels
                if paragraph.style.name.startswith('Heading'):
                    full_text.append(f"\n## {paragraph.text} ##\n")
                # Handle normal paragraphs
                else:
                    if paragraph.text.strip():  # Skip empty paragraphs
                        text = paragraph.text
                        # Add formatting indicators
                        if any(run.bold for run in paragraph.runs):
                            text = f"**{text}**"
                        if any(run.italic for run in paragraph.runs):
                            text = f"_{text}_"
                        full_text.append(text)
            
            # Handle tables if present
            for table in doc.tables:
                full_text.append("\nTable Contents:")
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    full_text.append(row_text)
            
            return '\n'.join(full_text)
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return f"Unsupported file type: {ext}"
    except Exception as e:
        return f"Error extracting text: {str(e)}"

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """Handle multiple file uploads and extract text"""
    results = []
    
    for file in files:
        # Validate file type
        allowed_extensions = {'.pdf', '.docx', '.txt'}
        file_ext = os.path.splitext(file.filename.lower())[1]
        
        if file_ext not in allowed_extensions:
            results.append({
                "filename": file.filename,
                "error": "Unsupported file type"
            })
            continue
            
        try:
            # Save file
            file_path = await save_uploaded_file(file)
            
            # Extract text
            extracted_text = extract_text(file_path)
            
            # Store document
            doc_id = document_store.store_document(
                filename=file.filename,
                file_type=file_ext,
                extracted_text=extracted_text
            )
            
            # Preprocess text
            tokens = preprocess_text(extracted_text)
            processed_text = tokens_to_string(tokens)
            
            # Update document with preprocessed text
            document_store.update_document(
                doc_id,
                preprocessed_text=processed_text
            )
            
            # Add to results
            results.append({
                "filename": file.filename,
                "doc_id": doc_id,
                "status": "success"
            })
            
        except Exception as e:
            results.append({
                "filename": file.filename,
                "error": str(e)
            })
    
    # After processing all documents, update vectors
    all_docs = document_store.get_all_documents()
    processed_texts = [
        doc['preprocessed_text']
        for doc in all_docs.values()
        if doc['preprocessed_text'] is not None
    ]
    
    if processed_texts:
        # Generate TF-IDF vectors
        vectors = document_vectorizer.fit_transform_documents(processed_texts)
        
        # Update documents with their vectors
        for (doc_id, doc), vector in zip(all_docs.items(), vectors):
            document_store.update_document(
                doc_id,
                vector=vector.toarray()[0].tolist()
            )
    
    return JSONResponse(content={
        "status": "success",
        "files": results
    })

@app.post("/cluster")
async def perform_clustering(num_clusters: int = 5):
    """
    Cluster all documents in the store
    Args:
        num_clusters: Number of clusters to create (default: 5)
    """
    # Get all documents
    docs = document_store.get_all_documents()
    if not docs:
        return JSONResponse(
            status_code=400,
            content={"message": "No documents available for clustering"}
        )
    
    # Get preprocessed texts and perform vectorization
    doc_ids = list(docs.keys())
    preprocessed_texts = [doc["preprocessed_text"] for doc in docs.values()]
    vectors = document_vectorizer.fit_transform_documents(preprocessed_texts)
    
    # Perform clustering
    labels, model = document_clusterer.cluster_documents(vectors, num_clusters)
    
    # Update document store with cluster labels
    for doc_id, label in zip(doc_ids, labels):
        document_store.update_document(doc_id, cluster=int(label))
    
    # Count documents per cluster
    unique_labels, counts = np.unique(labels, return_counts=True)
    cluster_distribution = {int(label): int(count) for label, count in zip(unique_labels, counts)}
    
    return {
        "message": "Clustering complete",
        "num_documents": len(doc_ids),
        "num_clusters": num_clusters,
        "cluster_distribution": cluster_distribution
    }

@app.get("/test-nltk")
async def test_nltk():
    """Test endpoint to verify NLTK functionality"""
    test_text = "Hello! This is a test sentence. We're testing NLTK processing."
    try:
        # Test preprocessing
        tokens = preprocess_text(test_text)
        processed = tokens_to_string(tokens)
        
        return {
            "status": "success",
            "original": test_text,
            "tokens": tokens,
            "processed": processed
        }
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }

@app.get("/")
def read_root():
    return {"Hello": "World"}

@app.get("/items/{item_id}")
def read_item(item_id: int, q: Union[str, None] = None):
    return {"item_id": item_id, "q": q}

@app.get("/document/{doc_id}")
async def get_document(doc_id: str):
    """Get document details including processed text"""
    doc = document_store.get_document(doc_id)
    if doc is None:
        return JSONResponse(
            status_code=404,
            content={"error": "Document not found"}
        )
    
    return {
        "filename": doc["filename"],
        "file_type": doc["file_type"],
        "extracted_text_preview": doc["extracted_text"][:200] + "...",  # Preview first 200 chars
        "preprocessed_text": doc["preprocessed_text"],
        "vector_length": len(doc["vector"]) if doc.get("vector") else 0
    }

@app.get("/documents")
async def list_documents():
    """List all documents and their preprocessing status"""
    docs = document_store.get_all_documents()
    return {
        "total_documents": len(docs),
        "documents": [
            {
                "doc_id": doc_id,
                "filename": doc["filename"],
                "status": "processed" if doc.get("vector") is not None else "pending"
            }
            for doc_id, doc in docs.items()
        ]
    }

@app.get("/cluster-contents")
async def get_cluster_contents():
    """Get documents grouped by their clusters"""
    return document_store.get_cluster_documents()

@app.delete("/document/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document from storage and file system"""
    try:
        # Get document before deletion to access filename
        doc = document_store.get_document(doc_id)
        if doc is None:
            return JSONResponse(
                status_code=404,
                content={"error": "Document not found"}
            )
        
        # Delete file from uploads folder
        file_path = os.path.join(UPLOAD_DIR, doc["filename"])
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Remove from document store
        document_store.delete_document(doc_id)
        
        return {"status": "success", "message": f"Document {doc_id} deleted"}
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to delete document: {str(e)}"}
        )

# Create semantic search instance
semantic_searcher = SemanticSearch()

@app.post("/semantic-search")
async def perform_semantic_search(query: str = Body(..., embed=True)):
    """
    Perform semantic search across documents
    Args:
        query: Search query string
    Returns:
        List of documents with similarity scores
    """
    # Get all documents
    docs = document_store.get_all_documents()
    if not docs:
        return JSONResponse(
            status_code=400,
            content={"message": "No documents available for search"}
        )
    
    # Update embeddings
    semantic_searcher.embed_documents(docs)
    
    # Perform search
    results = semantic_searcher.search(query)
    
    # Format results
    formatted_results = []
    for doc_id, similarity in results:
        doc = docs[doc_id]
        # Get a relevant snippet - first 200 chars for simplicity
        # In a real app, you'd want to do smart snippet extraction
        snippet = doc.get("extracted_text", "")[:200] + "..."
        
        formatted_results.append({
            "filename": doc["filename"],
            "snippet": snippet,
            "similarity": float(similarity),  # Convert to float for JSON serialization
            "doc_id": doc_id
        })
    
    return {
        "results": formatted_results
    }